# --- INICIO main.py v2.4.1-mt (Añade Provincia a Endpoint Place Details) ---
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from openai import OpenAI, APIError
import os
import shutil
import requests
import base64
import logging
import psycopg2  # Driver PostgreSQL
import psycopg2.extras  # Para DictCursor
import tempfile
import re
import chardet  # Necesario para extraer_texto_simple
import hashlib
import httpx  # Para llamadas async a Google Places API
from PyPDF2 import PdfReader, errors as pdf_errors
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# Configuración del Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler() # Asegura que los logs vayan a la consola/salida estándar
        # Podrías añadir logging.FileHandler('api.log') si quieres guardar en archivo
    ]
)
logger = logging.getLogger(__name__) # Usar un logger específico

# MODIFICADO: Versión indica adición endpoint Place Details y manejo provincia
app = FastAPI(
    title="Asistente IA UBIKUA API v2.4.1-mt (Place Details + Province)",
    version="2.4.1-mt",
    description="API para el Asistente IA UBIKUA con funcionalidades multi-tenant, RAG, y obtención de detalles de dirección."
)

# Configuración CORS (Considera restringir origins en producción)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # Cambiar a ["https://ia.ubikua.es", "https://ubikua.es"] en producción
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Configuración de Variables de Entorno ---
try:
    # OpenAI
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        logger.warning("Variable de entorno OPENAI_API_KEY no encontrada. Funcionalidad IA limitada.")
        client = None
    else:
        client = OpenAI(api_key=openai_api_key)
        logger.info("Cliente OpenAI configurado correctamente.")

    # Google Search (Custom Search API)
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY") # Para Google Search
    GOOGLE_CX = os.getenv("GOOGLE_CX")
    if not GOOGLE_API_KEY or not GOOGLE_CX:
        logger.warning("Variables GOOGLE_API_KEY o GOOGLE_CX (para Search) no encontradas. Búsqueda web deshabilitada.")
        SEARCH_CONFIGURED = False
    else:
        SEARCH_CONFIGURED = True
        logger.info("Credenciales Google Search API OK.")

    # Google Maps/Places API (para Place Details)
    MAPS_API_ALL = os.getenv("MAPS_API_ALL")
    if not MAPS_API_ALL:
        logger.warning("Variable MAPS_API_ALL no encontrada. Endpoint de detalles de dirección deshabilitado.")
        MAPS_CONFIGURED = False
    else:
        MAPS_CONFIGURED = True
        logger.info("Credencial MAPS_API_ALL (Places API) OK.")

    # Base de Datos PostgreSQL
    DB_HOST = os.getenv("DB_HOST")
    DB_USER = os.getenv("DB_USER")
    DB_PASS = os.getenv("DB_PASS")
    DB_NAME = os.getenv("DB_NAME")
    DB_PORT = int(os.getenv("DB_PORT", 5432))
    if not all([DB_HOST, DB_USER, DB_PASS, DB_NAME]):
        logger.warning("Faltan una o más variables de entorno para la BD PostgreSQL. Funcionalidad de BD deshabilitada.")
        DB_CONFIGURED = False
    else:
        DB_CONFIGURED = True
        logger.info("Credenciales BD PostgreSQL OK.")

    # Puente PHP (para servir documentos subidos)
    PHP_FILE_SERVE_URL = os.getenv("PHP_FILE_SERVE_URL")
    PHP_API_SECRET_KEY = os.getenv("PHP_API_SECRET_KEY")
    if not PHP_FILE_SERVE_URL or not PHP_API_SECRET_KEY:
        logger.warning("Faltan PHP_FILE_SERVE_URL o PHP_API_SECRET_KEY. Procesamiento de documentos (RAG) deshabilitado.")
        PHP_BRIDGE_CONFIGURED = False
    else:
        PHP_BRIDGE_CONFIGURED = True
        logger.info("Configuración PHP Bridge (File Serve) OK.")

except Exception as config_error:
    logger.error(f"Error crítico durante la carga de configuración inicial: {config_error}", exc_info=True)
    # Establecer flags a False para evitar operaciones si falla la configuración
    client = None
    SEARCH_CONFIGURED = False
    MAPS_CONFIGURED = False
    DB_CONFIGURED = False
    PHP_BRIDGE_CONFIGURED = False

# --- Modelos Pydantic ---
class PeticionConsulta(BaseModel):
    mensaje: str
    especializacion: str = "general"
    buscar_web: bool = False
    user_id: int | None = Field(None, description="ID del usuario que realiza la consulta")
    tenant_id: int | None = Field(None, description="ID del tenant/empresa del usuario")

class RespuestaConsulta(BaseModel):
    respuesta: str

class RespuestaAnalisis(BaseModel):
    informe: str

class ProcessRequest(BaseModel): # Modelo para /process-document
    doc_id: int = Field(..., description="ID del documento en la BD")
    user_id: int = Field(..., description="ID del usuario propietario")
    tenant_id: int | None = Field(None, description="ID del tenant/empresa propietario")

class ProcessResponse(BaseModel):
    success: bool
    message: str | None = None
    error: str | None = None

class PlaceDetailsResponse(BaseModel): # Modelo para /direccion/detalles
    success: bool
    postal_code: str | None = Field(None, description="Código Postal")
    locality: str | None = Field(None, description="Localidad / Ciudad")
    province: str | None = Field(None, description="Provincia / Estado / Región") # Provincia añadida
    country: str | None = Field(None, description="País")
    error: str | None = Field(None, description="Mensaje de error si success es false")

# --- Prompts (Sin cambios respecto a versiones anteriores) ---
BASE_PROMPT_CONSULTA = (
    "Eres el Asistente IA oficial de UBIKUA, un experto multidisciplinar que responde de manera "
    "clara, precisa y estéticamente impecable. Tus respuestas deben estar redactadas en HTML válido, "
    "usando etiquetas como <h2> y <h3> para encabezados, <p> para párrafos, <strong> para destacar información, "
    "y <ul> o <ol> para listas cuando sea necesario. Además, si la información lo requiere, utiliza etiquetas "
    "de tabla (<table>, <thead>, <tbody>, <tr>, <td>) para organizar datos de manera clara y estructurada. "
    "Debes evitar el uso de markdown o estilo plano. Asegúrate de que cada respuesta tenga una estructura lógica: "
    "comienza con un título principal, seguido de secciones bien delimitadas, tablas (si corresponde) y una "
    "conclusión clara. Siempre utiliza un tono profesional y adaptado al usuario, ofreciendo ejemplos y resúmenes "
    "que garanticen una comprensión total del contenido. Por favor, responde solo con HTML sin ningún comentario "
    "de código o markdown adicional."
)
BASE_PROMPT_ANALISIS_DOC = (
    "Eres el Asistente IA oficial de UBIKUA y un experto en redactar informes y análisis de documentos. "
    "A partir del texto suministrado, redacta un informe completo, profesional y estéticamente bien estructurado en HTML. "
    "Utiliza etiquetas HTML como <h1>, <h2> para títulos y subtítulos, <p> para párrafos, <strong> para resaltar puntos clave, "
    "y <table> con <thead>, <tbody>, <tr> y <td> para presentar datos o resúmenes numéricos. La respuesta debe incluir: "
    "1) un título principal, 2) secciones con encabezados relevantes, 3) listas o tablas cuando corresponda, y 4) una conclusión. "
    "Asegúrate de que la salida final sea un documento que se pueda copiar y pegar en Word o Google Docs sin perder el formato."
    "Por favor, responde solo con HTML sin ningún comentario de código o markdown adicional."
)
PROMPT_ESPECIALIZACIONES = {
    "general": "Ofrece una respuesta amplia, comprensiva y detallada, abarcando todos los puntos relevantes de la consulta.",
    "legal": "Adopta un enfoque riguroso y formal, utilizando terminología jurídica adecuada y estructurando la respuesta de forma clara y precisa.",
    "comunicacion": "Emplea un estilo persuasivo y creativo, con ejemplos y metáforas que faciliten la comprensión, y asegura que la respuesta sea atractiva y comunicativa.",
    "formacion": "Proporciona explicaciones didácticas y detalladas, estructuradas en secciones claramente delimitadas, con ejemplos prácticos y casos ilustrativos para facilitar el aprendizaje.",
    "informatica": "Ofrece respuestas técnicas precisas, explicando conceptos y procesos de tecnología de forma clara, con ejemplos, pseudocódigo o diagramas si es necesario.",
    "direccion": "Brinda una perspectiva estratégica, analizando tendencias y ofreciendo recomendaciones ejecutivas y bien fundamentadas, siempre con un tono profesional y asertivo.",
    "innovacion": "Responde de manera creativa e innovadora, proponiendo ideas disruptivas y soluciones fuera de lo convencional, utilizando analogías y ejemplos que inspiren nuevas perspectivas.",
    "contabilidad": "Proporciona respuestas precisas y estructuradas, con terminología contable adecuada, apoyadas en ejemplos numéricos y análisis detallados cuando sea relevante.",
    "administracion": "Enfócate en la eficiencia y organización, ofreciendo análisis claros sobre procesos, recomendaciones prácticas y estructuradas que faciliten la gestión y toma de decisiones."
}
FRASES_BUSQUEDA = ["no tengo información", "no dispongo de información", "no tengo acceso", "no sé"]

# --- Directorio Temporal para Subidas ---
# Render provee un sistema de archivos efímero en /tmp
TEMP_DIR = "/tmp/ubikua_uploads"
try:
    os.makedirs(TEMP_DIR, exist_ok=True)
    logger.info(f"Directorio temporal para subidas verificado/creado en: {TEMP_DIR}")
except OSError as e:
    logger.error(f"No se pudo crear el directorio temporal {TEMP_DIR}: {e}. El procesamiento de archivos fallará.")
    # Considerar si la API debe arrancar si esto falla

# --- Funciones Auxiliares ---

def get_db_connection():
    """Establece y devuelve una conexión a la base de datos PostgreSQL."""
    if not DB_CONFIGURED:
        logger.error("Configuración de base de datos incompleta. No se puede conectar.")
        return None
    try:
        conn = psycopg2.connect(
            host=DB_HOST, database=DB_NAME, user=DB_USER,
            password=DB_PASS, port=DB_PORT, connect_timeout=5 # Timeout 5 seg
        )
        # logger.debug("Conexión PostgreSQL establecida con éxito.") # Log más verboso
        return conn
    except (Exception, psycopg2.Error) as error:
        logger.error(f"Error al conectar con PostgreSQL: {error}", exc_info=True)
        return None

def extraer_texto_pdf_docx(ruta_archivo: str, extension: str) -> str:
    """Extrae texto de archivos PDF, DOC y DOCX."""
    texto = ""
    logger.info(f"Extrayendo texto ({extension.upper()}) de: {os.path.basename(ruta_archivo)}")
    try:
        if extension == "pdf":
            with open(ruta_archivo, 'rb') as archivo:
                lector = PdfReader(archivo, strict=False) # strict=False para ser más tolerante
                if lector.is_encrypted:
                    logger.warning(f"El archivo PDF '{os.path.basename(ruta_archivo)}' está encriptado. La extracción puede fallar.")
                    # Podría intentar desbloquear con contraseña vacía: lector.decrypt('')
                num_paginas = len(lector.pages)
                logger.info(f"Procesando {num_paginas} páginas del PDF.")
                for i, pagina in enumerate(lector.pages):
                    try:
                        texto_pagina = pagina.extract_text()
                        if texto_pagina:
                            texto += texto_pagina + "\n"
                    except Exception as page_error:
                        logger.warning(f"Error al extraer texto de página {i+1} en {os.path.basename(ruta_archivo)}: {page_error}")
                        continue # Continuar con la siguiente página
        elif extension in ["doc", "docx"]:
             try:
                doc = Document(ruta_archivo)
                parrafos_texto = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
                texto = "\n".join(parrafos_texto)
                # Podríamos añadir extracción de tablas si es necesario
                # for table in doc.tables:
                #    for row in table.rows:
                #        for cell in row.cells:
                #            texto += cell.text + "\t" # Ejemplo simple
                #        texto += "\n"

             except PackageNotFoundError:
                 logger.error(f"Error al leer DOCX '{os.path.basename(ruta_archivo)}': Archivo no encontrado o formato inválido.")
                 return "[Error DOCX: Archivo inválido o no encontrado]"
             except Exception as docx_error: # Captura genérica para otros errores de python-docx
                 logger.error(f"Error inesperado al procesar DOCX '{os.path.basename(ruta_archivo)}': {docx_error}", exc_info=True)
                 return "[Error interno procesando DOCX]"
        else:
            logger.error(f"Tipo de archivo no esperado '{extension}' en extraer_texto_pdf_docx")
            return "[Error interno: Tipo no esperado]"

        texto_limpio = texto.strip()
        if not texto_limpio:
             logger.warning(f"No se extrajo texto útil (vacío después de strip) de '{os.path.basename(ruta_archivo)}'")
             return "[Archivo sin texto extraíble]"

        logger.info(f"Texto extraído de {extension.upper()} OK (longitud: {len(texto_limpio)} caracteres).")
        return texto_limpio

    except pdf_errors.PdfReadError as e:
        logger.error(f"Error crítico al leer PDF '{os.path.basename(ruta_archivo)}': {e}")
        return "[Error PDF: No se pudo leer o archivo corrupto]"
    except FileNotFoundError:
        logger.error(f"Error interno de sistema: Archivo '{os.path.basename(ruta_archivo)}' no encontrado durante la extracción.")
        return "[Error interno: Archivo no encontrado]"
    except Exception as e:
        logger.error(f"Error general al extraer texto de {extension.upper()} '{os.path.basename(ruta_archivo)}': {e}", exc_info=True)
        return f"[Error interno procesando {extension.upper()}]"


def extraer_texto_simple(ruta_archivo: str) -> str:
    """Extrae texto de archivos planos (TXT, CSV) detectando encoding."""
    logger.info(f"Extrayendo texto simple de: {os.path.basename(ruta_archivo)}")
    texto = ""
    detected_encoding = 'utf-8' # Default
    try:
        with open(ruta_archivo, 'rb') as fb:
            raw_data = fb.read()
            if not raw_data:
                logger.warning(f"Archivo vacío: {os.path.basename(ruta_archivo)}")
                return ""
            # Detectar encoding
            detection = chardet.detect(raw_data)
            detected_encoding = detection['encoding'] if detection['encoding'] and detection['confidence'] > 0.5 else 'utf-8'
            if not detection['encoding'] or detection['confidence'] <= 0.5:
                 logger.info(f"Detección de encoding incierta ({detection}), usando '{detected_encoding}'.")
            else:
                 logger.info(f"Encoding detectado: {detected_encoding} (Confianza: {detection['confidence']:.2f})")

        with open(ruta_archivo, 'r', encoding=detected_encoding, errors='ignore') as f:
            texto = f.read()

        texto_limpio = texto.strip()
        if not texto_limpio:
             logger.warning(f"No se extrajo texto útil (vacío después de strip) de '{os.path.basename(ruta_archivo)}'")
             return "[Archivo sin texto extraíble]"

        logger.info(f"Texto extraído simple OK (longitud: {len(texto_limpio)} caracteres).")
        return texto_limpio

    except FileNotFoundError:
        logger.error(f"Error interno: Archivo '{os.path.basename(ruta_archivo)}' no encontrado para extracción simple.")
        return "[Error interno: Archivo no encontrado]"
    except UnicodeDecodeError as ude:
         logger.error(f"Error de decodificación para '{os.path.basename(ruta_archivo)}' con encoding '{detected_encoding}': {ude}")
         # Intento con otro encoding común como fallback
         try:
             logger.info(f"Intentando decodificar '{os.path.basename(ruta_archivo)}' con ISO-8859-1...")
             with open(ruta_archivo, 'r', encoding='iso-8859-1', errors='ignore') as f_fallback:
                 texto = f_fallback.read()
             texto_limpio = texto.strip()
             if not texto_limpio: return "[Archivo sin texto extraíble]"
             logger.info(f"Texto extraído simple con fallback ISO-8859-1 OK (longitud: {len(texto_limpio)} caracteres).")
             return texto_limpio
         except Exception as e_fallback:
             logger.error(f"Fallo el fallback de decodificación para '{os.path.basename(ruta_archivo)}': {e_fallback}")
             return "[Error de Codificación: No se pudo leer el texto]"
    except Exception as e:
        logger.error(f"Error inesperado al extraer texto simple de '{os.path.basename(ruta_archivo)}': {e}", exc_info=True)
        return "[Error interno procesando texto plano]"


def buscar_google(query: str) -> str:
    """Realiza una búsqueda en Google Custom Search y devuelve resultados formateados en HTML."""
    if not SEARCH_CONFIGURED:
        logger.warning("Llamada a buscar_google pero SEARCH_CONFIGURED es False.")
        return "<p><i>[Búsqueda web no disponible (sin configurar).]</i></p>"

    url = "https://www.googleapis.com/customsearch/v1"
    params = {"key": GOOGLE_API_KEY, "cx": GOOGLE_CX, "q": query, "num": 3, "lr": "lang_es"} # Limitar a 3 resultados y preferir español
    logger.info(f"Buscando en Google (Custom Search): '{query}'")
    try:
        # Usar requests aquí está bien, ya que buscar_google se llama desde un endpoint síncrono (/consulta)
        response = requests.get(url, params=params, timeout=10)
        response.raise_for_status() # Lanza excepción para errores HTTP
        data = response.json()

        # Comprobar si hay errores específicos de la API de Google Search
        if "error" in data:
             error_details = data["error"].get("message", "Error desconocido de Google Search API")
             error_code = data["error"].get("code", "N/A")
             logger.error(f"Error de Google Search API: {error_details} (Código: {error_code})")
             return f"<p><i>[Error en la búsqueda web: {error_details}]</i></p>"

        resultados = data.get("items", [])
        if not resultados:
            logger.info("Búsqueda web no devolvió resultados.")
            return "<p><i>[No se encontraron resultados web relevantes.]</i></p>"

        # Formatear resultados como lista HTML
        texto_resultados = "<div class='google-results' style='margin-top:15px; padding-top:10px; border-top: 1px solid #eee;'>"
        texto_resultados += "<h4 style='font-size:0.9em;color:#555; margin-bottom: 5px;'>Resultados web relacionados:</h4><ul>"
        for item in resultados:
            title = item.get('title','Sin título')
            link = item.get('link','#')
            # Limpiar snippet para evitar HTML inyectado y saltos de línea
            snippet = item.get('snippet','')
            if snippet:
                 snippet = re.sub('<.*?>', '', snippet) # Quitar tags HTML del snippet
                 snippet = snippet.replace('\n',' ').strip()
            else:
                 snippet = "No hay descripción disponible."

            texto_resultados += (
                f"<li style='margin-bottom: 10px; padding-left: 5px; border-left: 3px solid #ddd;'>"
                f"<a href='{link}' target='_blank' style='font-weight: bold; color: #1a0dab; text-decoration: none; display: block; margin-bottom: 2px;'>{htmlspecialchars(title)}</a>"
                f"<p style='font-size: 0.85em; margin: 0; color: #333;'>{htmlspecialchars(snippet)}</p>"
                f"<cite style='font-size: 0.8em; color: #006621; display: block; margin-top: 2px;'>{htmlspecialchars(link)}</cite>"
                f"</li>"
            )
        texto_resultados += "</ul></div>"
        logger.info(f"Búsqueda web OK: {len(resultados)} resultados formateados.")
        return texto_resultados

    except requests.exceptions.Timeout:
        logger.error("Timeout durante la búsqueda web.")
        return "<p><i>[Error: La búsqueda web tardó demasiado en responder.]</i></p>"
    except requests.exceptions.RequestException as e:
        logger.error(f"Error en la solicitud de búsqueda web: {e}")
        return "<p><i>[Error de conexión durante la búsqueda web.]</i></p>"
    except Exception as e:
        logger.error(f"Error inesperado durante la búsqueda web: {e}", exc_info=True)
        return "<p><i>[Error inesperado durante la búsqueda web.]</i></p>"

# Helper para escapar HTML (importante para mostrar snippets de forma segura)
from html import escape as htmlspecialchars

# --- Endpoint /process-document ---
@app.post("/process-document", response_model=ProcessResponse)
async def process_document_text(request: ProcessRequest):
    """
    Procesa un documento previamente subido: obtiene su contenido (vía PHP bridge),
    extrae el texto y lo guarda en la base de datos junto con la marca 'procesado'.
    También actualiza el vector FTS (manejado por trigger en BD).
    """
    doc_id = request.doc_id
    current_user_id = request.user_id
    current_tenant_id = request.tenant_id

    # Validar IDs de entrada
    if not isinstance(current_user_id, int) or not isinstance(current_tenant_id, int):
        logger.error(f"IDs inválidos recibidos en /process-document: User='{current_user_id}', Tenant='{current_tenant_id}' para Doc ID {doc_id}")
        # Devolver 400 Bad Request si los IDs no son enteros válidos
        return ProcessResponse(success=False, error="User ID y Tenant ID deben ser números enteros válidos.")

    logger.info(f"Iniciando procesamiento para Documento ID: {doc_id} (Usuario: {current_user_id}, Tenant: {current_tenant_id})")

    # Verificar configuración necesaria
    if not DB_CONFIGURED or not PHP_BRIDGE_CONFIGURED:
        error_msg = "Configuración incompleta en el backend para procesar documentos."
        if not DB_CONFIGURED: error_msg += " (Falta config BD)"
        if not PHP_BRIDGE_CONFIGURED: error_msg += " (Falta config PHP Bridge)"
        logger.error(error_msg + f" - Solicitud para doc {doc_id}")
        return ProcessResponse(success=False, error=error_msg)

    conn = None
    temp_path = None # Path al archivo temporal

    try:
        # 1. Obtener información del documento desde la BD
        conn = get_db_connection()
        if not conn:
            # Error ya logueado en get_db_connection
            raise ConnectionError("No se pudo establecer conexión con la BD.")

        original_fname = None
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            logger.debug(f"Consultando BD info para doc {doc_id}, user {current_user_id}, tenant {current_tenant_id}")
            sql_select = """
                SELECT original_filename, file_type, stored_path, procesado
                FROM user_documents
                WHERE id = %s AND user_id = %s AND tenant_id = %s
            """
            cursor.execute(sql_select, (doc_id, current_user_id, current_tenant_id))
            doc_info = cursor.fetchone()

            if not doc_info:
                logger.error(f"Documento ID {doc_id} no encontrado en BD para User {current_user_id}, Tenant {current_tenant_id}.")
                raise FileNotFoundError(f"Documento ID {doc_id} no encontrado para este usuario/tenant.")

            original_fname = doc_info['original_filename']
            file_type = doc_info['file_type']
            stored_path = doc_info['stored_path'] # No usado si usamos PHP Bridge
            is_already_processed = doc_info['procesado']

            if is_already_processed:
                 logger.info(f"Documento {doc_id} ('{original_fname}') ya estaba marcado como procesado. Omitiendo re-procesamiento.")
                 # Considerar si devolver éxito o un mensaje específico. Éxito parece razonable.
                 return ProcessResponse(success=True, message="El documento ya estaba procesado.")

            logger.info(f"Info obtenida para doc {doc_id}: fname='{original_fname}', type='{file_type}'")

        # Cerrar conexión temprana si ya no se necesita para la obtención de datos
        conn.close()
        conn = None # Marcar como cerrada

        # 2. Obtener contenido del archivo vía PHP Bridge
        serve_url = f"{PHP_FILE_SERVE_URL}?doc_id={doc_id}&user_id={current_user_id}&tenant_id={current_tenant_id}&api_key={PHP_API_SECRET_KEY}"
        logger.info(f"Solicitando contenido de doc ID {doc_id} al PHP Bridge. URL: {serve_url}")

        # Usamos requests síncrono por simplicidad aquí, podría ser async con httpx si esto se vuelve un cuello de botella
        response_full = requests.get(serve_url, timeout=120, stream=True) # Timeout más largo para archivos grandes
        response_full.raise_for_status() # Lanza excepción para códigos 4xx/5xx
        logger.info(f"Respuesta recibida de PHP Bridge (Status: {response_full.status_code}). Iniciando descarga a temporal...")

        # 3. Guardar en archivo temporal y extraer texto
        file_ext = os.path.splitext(original_fname)[1].lower().strip('.') if original_fname else ''
        extracted_text = None
        TEXT_EXTENSIONS_PROC = ["pdf", "doc", "docx", "txt", "csv"] # Extensiones que procesamos

        if file_ext in TEXT_EXTENSIONS_PROC:
            # Crear archivo temporal de forma segura usando NamedTemporaryFile
            with tempfile.NamedTemporaryFile(mode='wb', suffix=f'.{file_ext}', dir=TEMP_DIR, delete=False) as temp_file:
                temp_path = temp_file.name
                bytes_written = 0
                try:
                    logger.info(f"Guardando contenido en archivo temporal: {temp_path}")
                    for chunk in response_full.iter_content(chunk_size=8192):
                        temp_file.write(chunk)
                        bytes_written += len(chunk)
                    logger.info(f"Escritos {bytes_written} bytes en {temp_path}")
                    if bytes_written == 0:
                         logger.warning(f"El archivo recibido de PHP Bridge para doc {doc_id} parece estar vacío.")
                         extracted_text = "[Archivo vacío recibido]" # Marcar como vacío

                except Exception as write_err:
                     logger.error(f"Error escribiendo en archivo temporal {temp_path}: {write_err}", exc_info=True)
                     raise IOError(f"No se pudo escribir el archivo temporal para {original_fname}")

            # Extraer texto solo si el archivo no estaba vacío
            if extracted_text is None: # Si no se marcó como vacío
                 logger.info(f"Archivo temporal {temp_path} escrito y cerrado. Extrayendo texto...")
                 if file_ext in ['pdf', 'doc', 'docx']:
                     extracted_text = extraer_texto_pdf_docx(temp_path, file_ext)
                 elif file_ext in ['txt', 'csv']:
                     extracted_text = extraer_texto_simple(temp_path)
                 logger.info(f"Texto extraído para doc {doc_id} (longitud: {len(extracted_text) if extracted_text else 0} caracteres)")
        else:
             logger.warning(f"Extracción no soportada para la extensión '{file_ext}' del archivo {original_fname} (doc {doc_id})")
             extracted_text = f"[Extracción no soportada para tipo: {file_ext}]"

        # Limpieza final del archivo temporal (fuera del 'with' para asegurar que está cerrado)
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                logger.info(f"Archivo temporal {temp_path} eliminado tras extracción.")
            except OSError as e_remove:
                logger.error(f"Error al borrar archivo temporal {temp_path} tras extracción: {e_remove}")
            temp_path = None # Asegurar que no se intente borrar de nuevo en finally

        # Validar texto extraído antes de actualizar BD
        if extracted_text is None or extracted_text.strip() == "" or extracted_text.strip().startswith("[Error"):
             if extracted_text is None: extracted_text = "[Error: Extracción devolvió None]"
             if not extracted_text.strip(): extracted_text = "[Archivo vacío o sin texto extraíble]"
             logger.error(f"Extracción de texto fallida o vacía para doc {doc_id}. Contenido guardado en BD: '{extracted_text[:100]}...'")
             # Guardaremos este 'error' o mensaje en la BD para indicar el problema

        # 4. Actualizar la base de datos marcando como procesado y guardando el texto
        logger.info(f"Actualizando BD para doc ID {doc_id} / tenant {current_tenant_id}...")
        conn = get_db_connection() # Re-abrir conexión para el update
        if not conn:
             raise ConnectionError("No se pudo reconectar a la BD para guardar el texto extraído.")

        with conn.cursor() as cursor:
            sql_update = """
                UPDATE user_documents
                SET extracted_text = %s, procesado = TRUE
                WHERE id = %s AND user_id = %s AND tenant_id = %s
            """
            # Truncar texto si excede un límite razonable (ej. 10MB de texto plano)
            MAX_TEXT_LENGTH = 10 * 1024 * 1024
            if len(extracted_text) > MAX_TEXT_LENGTH:
                 logger.warning(f"Texto extraído truncado a {MAX_TEXT_LENGTH} caracteres para BD (doc {doc_id}). Longitud original: {len(extracted_text)}")
                 extracted_text_to_save = extracted_text[:MAX_TEXT_LENGTH]
            else:
                 extracted_text_to_save = extracted_text

            cursor.execute(sql_update, (extracted_text_to_save, doc_id, current_user_id, current_tenant_id))
            rows_affected = cursor.rowcount
            conn.commit() # Hacer commit de la transacción

            if rows_affected == 0:
                # Esto es preocupante si el SELECT inicial lo encontró. Podría indicar un problema.
                logger.warning(f"UPDATE no afectó filas al guardar texto del doc {doc_id}/Tenant {current_tenant_id}. ¿El documento fue modificado/eliminado concurrentemente?")
                # Considerar devolver error si esto no debería pasar. Por ahora, continuamos pero logueamos.
            else:
                 logger.info(f"Base de datos actualizada con éxito para doc ID {doc_id} ({rows_affected} fila afectada).")

        return ProcessResponse(success=True, message="Documento procesado y texto extraído.")

    # --- Bloques de Manejo de Excepciones ---
    except FileNotFoundError as e:
         logger.error(f"Error FNF procesando doc {doc_id}: {e}")
         return ProcessResponse(success=False, error=str(e)) # Error específico para el frontend
    except ConnectionError as e: # Nuestro error de conexión a BD
         # Error ya logueado en get_db_connection o en los reintentos
         return ProcessResponse(success=False, error=f"Error de conexión a la base de datos: {e}")
    except requests.exceptions.RequestException as e:
        logger.error(f"Error al conectar con PHP Bridge para doc {doc_id}: {e}", exc_info=True)
        status_code = e.response.status_code if e.response is not None else 'N/A'
        return ProcessResponse(success=False, error=f"Error al obtener el archivo ({status_code}). Verifica el servicio de archivos.")
    except IOError as e: # Error de escritura/lectura de archivo temporal
         logger.error(f"Error de I/O con archivo temporal para doc {doc_id}: {e}", exc_info=True)
         return ProcessResponse(success=False, error=f"Error al manejar archivo temporal: {e}")
    except (psycopg2.Error) as db_err: # Errores específicos de PostgreSQL
        logger.error(f"Error de Base de Datos procesando doc {doc_id}: {db_err}", exc_info=True)
        if conn and not conn.closed: conn.rollback() # Revertir si hubo error de BD
        return ProcessResponse(success=False, error=f"Error de base de datos durante el procesamiento.")
    except Exception as e: # Captura genérica para otros errores inesperados
        logger.error(f"Error general inesperado procesando doc {doc_id}: {e}", exc_info=True)
        return ProcessResponse(success=False, error=f"Error interno del servidor ({type(e).__name__}). Contacte al administrador.")
    finally:
        # Limpieza final del archivo temporal si por alguna razón aún existe
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                logger.info(f"Archivo temporal {temp_path} eliminado en bloque finally (verificación).")
            except OSError as e_remove:
                logger.error(f"Error al borrar archivo temporal {temp_path} en finally: {e_remove}")
        # Asegurar que la conexión a BD esté cerrada
        if conn and not conn.closed:
            conn.close()
            # logger.debug("Conexión a BD cerrada en finally (process-document).")


# --- Endpoint de consulta (/consulta) ---
@app.post("/consulta", response_model=RespuestaConsulta)
def consultar_agente(datos: PeticionConsulta):
    """
    Procesa una consulta del usuario: construye el prompt (base + especialización + memoria + RAG),
    llama a OpenAI, realiza búsqueda web si es necesario y devuelve la respuesta formateada en HTML.
    """
    # Verificar configuración OpenAI
    if not client:
        logger.error("Llamada a /consulta pero cliente OpenAI no está configurado.")
        raise HTTPException(status_code=503, detail="Servicio IA no disponible (Configuración).")

    # Validar IDs de entrada
    current_user_id = datos.user_id
    current_tenant_id = datos.tenant_id
    if not isinstance(current_user_id, int) or not isinstance(current_tenant_id, int):
        logger.error(f"IDs inválidos recibidos en /consulta: User='{current_user_id}', Tenant='{current_tenant_id}'")
        raise HTTPException(status_code=400, detail="User ID y Tenant ID deben ser números enteros válidos.")

    # Preparar datos de entrada
    especializacion = datos.especializacion.lower() if datos.especializacion else "general"
    mensaje_usuario = datos.mensaje.strip()
    forzar_busqueda_web = datos.buscar_web
    logger.info(f"Consulta recibida: User={current_user_id}, Tenant={current_tenant_id}, Espec='{especializacion}', WebForzado={forzar_busqueda_web}, Msg='{mensaje_usuario[:100]}...'")

    if not mensaje_usuario:
         logger.warning("Mensaje de usuario vacío recibido en /consulta.")
         # Devolver una respuesta indicando que se necesita un mensaje
         return RespuestaConsulta(respuesta="<p>Por favor, introduce tu consulta.</p>")

    # --- Obtener prompt personalizado (Memoria) ---
    custom_prompt_text = ""
    if DB_CONFIGURED:
        conn_prompt = get_db_connection()
        if conn_prompt:
            try:
                with conn_prompt.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
                    sql_prompt = "SELECT custom_prompt FROM user_settings WHERE user_id = %s AND tenant_id = %s"
                    cursor.execute(sql_prompt, (current_user_id, current_tenant_id))
                    result = cursor.fetchone()
                    if result and result.get('custom_prompt') and result['custom_prompt'].strip():
                        custom_prompt_text = result['custom_prompt'].strip()
                        logger.info(f"Prompt personalizado (Memoria) OK para user {current_user_id}/tenant {current_tenant_id}.")
                    else:
                         logger.info(f"No hay prompt personalizado (Memoria) para user {current_user_id}/tenant {current_tenant_id}.")
            except (Exception, psycopg2.Error) as e:
                logger.error(f"Error BD get prompt (Memoria) user {current_user_id}/tenant {current_tenant_id}: {e}", exc_info=True)
            finally:
                conn_prompt.close()
        else:
            logger.warning(f"No conexión BD para prompt (Memoria) user {current_user_id}")

    # --- Obtener contexto documental (RAG) ---
    document_context = ""
    MAX_RAG_TOKENS = 3500 # Límite de tokens para contexto RAG
    if DB_CONFIGURED:
        conn_docs = get_db_connection()
        if conn_docs:
            try:
                with conn_docs.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
                    # Preparar query FTS de forma segura
                    search_query_cleaned = re.sub(r'[!\'()|&:*<>~@]', ' ', mensaje_usuario).strip() # Quitar más caracteres problemáticos para FTS
                    search_query_terms = search_query_cleaned.split()
                    if not search_query_terms:
                         logger.info(f"Mensaje de usuario para RAG vacío o sin términos válidos tras limpiar.")
                    else:
                        # Crear query tsquery: term1 & term2 & ...
                        fts_query_string = ' & '.join(search_query_terms)
                        logger.info(f"Ejecutando búsqueda RAG FTS con query: '{fts_query_string}' para user {current_user_id}/tenant {current_tenant_id}")

                        sql_fts = """
                            SELECT original_filename, extracted_text,
                                   ts_rank_cd(fts_vector, plainto_tsquery('spanish', %(query)s)) as relevance
                            FROM user_documents
                            WHERE user_id = %(user_id)s AND tenant_id = %(tenant_id)s
                              AND is_active_for_ai = TRUE
                              AND procesado = TRUE -- Solo usar documentos procesados
                              AND fts_vector @@ plainto_tsquery('spanish', %(query)s)
                              AND extracted_text IS NOT NULL AND extracted_text != '' AND NOT extracted_text LIKE '[Error%%' AND NOT extracted_text LIKE '[Archivo vacío%%'
                            ORDER BY relevance DESC
                            LIMIT 5 -- Obtener los 5 más relevantes y filtrar por tokens
                        """
                        cursor.execute(sql_fts, {'query': fts_query_string, 'user_id': current_user_id, 'tenant_id': current_tenant_id})
                        relevant_docs = cursor.fetchall()

                        if relevant_docs:
                            logger.info(f"Encontrados {len(relevant_docs)} documentos RAG potenciales (pre-filtrado tokens).")
                            context_parts = ["\n\n### Contexto de tus Documentos ###\n(Información extraída de tus archivos para ayudar a responder tu consulta)\n"]
                            current_token_count = 0
                            docs_included_count = 0

                            for doc in relevant_docs:
                                filename = doc['original_filename']
                                text = doc['extracted_text']
                                relevance_score = doc['relevance']
                                logger.debug(f"Evaluando doc RAG: '{filename}' (Relevancia: {relevance_score:.4f})")

                                # Estimación simple de tokens (considerar tiktoken para precisión)
                                doc_tokens_estimated = len(text.split()) * 1.3

                                if current_token_count + doc_tokens_estimated <= MAX_RAG_TOKENS:
                                    context_parts.append(f"\n--- Documento: {htmlspecialchars(filename)} (Relevancia: {relevance_score:.2f}) ---")
                                    context_parts.append(text) # Asumimos que el texto ya está limpio
                                    # context_parts.append(f"--- Fin Documento: {htmlspecialchars(filename)} ---") # Opcional
                                    current_token_count += doc_tokens_estimated
                                    docs_included_count += 1
                                    logger.debug(f"Añadido doc '{filename}'. Tokens acumulados: ~{current_token_count:.0f}/{MAX_RAG_TOKENS}")
                                    if current_token_count >= MAX_RAG_TOKENS:
                                         logger.warning(f"Alcanzado límite RAG ({MAX_RAG_TOKENS} tokens) tras incluir {docs_included_count} docs.")
                                         break
                                else:
                                    # Intentar añadir solo una porción si cabe algo significativo
                                    remaining_tokens = MAX_RAG_TOKENS - current_token_count
                                    # Definir un umbral mínimo de tokens para considerar añadir una porción
                                    MIN_PARTIAL_TOKENS = 150
                                    if remaining_tokens > MIN_PARTIAL_TOKENS:
                                         # Calcular caracteres basados en tokens restantes
                                         available_chars = max(100, int(remaining_tokens / 1.3)) # Añadir al menos 100 chars
                                         partial_text = text[:available_chars]
                                         context_parts.append(f"\n--- Documento (Parcial): {htmlspecialchars(filename)} (Relevancia: {relevance_score:.2f}) ---")
                                         context_parts.append(partial_text + "...") # Indicar que está truncado
                                         # context_parts.append(f"--- Fin Documento (Parcial): {htmlspecialchars(filename)} ---") # Opcional
                                         current_token_count += remaining_tokens # Llenar el cupo
                                         docs_included_count += 1
                                         logger.warning(f"Incluida porción parcial de '{filename}'. Alcanzado límite RAG ({MAX_RAG_TOKENS} tokens).")
                                    else:
                                         logger.info(f"Doc '{filename}' omitido, no caben suficientes tokens restantes ({remaining_tokens:.0f} < {MIN_PARTIAL_TOKENS}).")
                                    break # Salir tras añadir porción o si no cabe nada

                            if docs_included_count > 0:
                               document_context = "\n".join(context_parts)
                               logger.info(f"Contexto RAG construido con {docs_included_count} documentos (~{current_token_count:.0f} tokens).")
                            else:
                                 logger.info("Ningún documento RAG añadido al contexto final (límite tokens o relevancia).")
                        else:
                            logger.info(f"No se encontraron documentos RAG para la consulta.")
            except (Exception, psycopg2.Error) as e:
                logger.error(f"Error al obtener contexto RAG de BD para user {current_user_id}/tenant {current_tenant_id}: {e}", exc_info=True)
                # No fallar toda la consulta por error RAG, continuar sin contexto documental
                document_context = "\n\n<p><i>[Nota: Hubo un problema al buscar información en tus documentos.]</i></p>"
            finally:
                if conn_docs:
                    conn_docs.close()
        else:
            logger.warning(f"No conexión BD para RAG user {current_user_id}")

    # --- Combinar prompts ---
    prompt_especifico = PROMPT_ESPECIALIZACIONES.get(especializacion, PROMPT_ESPECIALIZACIONES["general"])
    system_prompt_parts = [BASE_PROMPT_CONSULTA, prompt_especifico]
    if custom_prompt_text:
        system_prompt_parts.append(f"\n\n### Instrucciones Adicionales (Memoria) ###\n{custom_prompt_text}")
    if document_context:
        system_prompt_parts.append(document_context) # Ya tiene su propio encabezado

    system_prompt = "\n".join(filter(None, system_prompt_parts))
    # logger.debug(f"System Prompt Final para OpenAI:\n------\n{system_prompt}\n------") # Descomentar para depuración profunda

    # --- Llamada a OpenAI ---
    texto_respuesta_final = "<p><i>Error inesperado al generar la respuesta.</i></p>" # Default error response
    MAX_RETRIES_OPENAI = 2
    for attempt in range(MAX_RETRIES_OPENAI):
         try:
            logger.info(f"Llamando a OpenAI (Intento {attempt + 1}/{MAX_RETRIES_OPENAI}) para User {current_user_id}/Tenant {current_tenant_id}...")
            # Medir tiempo de llamada OpenAI?
            # start_time = time.time()
            respuesta_inicial = client.chat.completions.create(
                model="gpt-4-turbo", # Asegúrate que este modelo está disponible y es el adecuado
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": mensaje_usuario}
                ],
                temperature=0.6, # Ligeramente más creativo si hay RAG/Memoria
                max_tokens=2000 # Límite de tokens en la respuesta
                # Podría añadirse user=f"user_{current_user_id}_tenant_{current_tenant_id}" para monitorización OpenAI
            )
            # elapsed_time = time.time() - start_time
            # logger.info(f"Respuesta OpenAI recibida en {elapsed_time:.2f} segundos.")

            if not respuesta_inicial.choices or not respuesta_inicial.choices[0].message or not respuesta_inicial.choices[0].message.content:
                 logger.error("Respuesta de OpenAI inválida o vacía.")
                 texto_respuesta_final = "<p><i>Error: No se recibió una respuesta válida del servicio de IA.</i></p>"
                 # Considerar si reintentar o fallar directamente
                 if attempt < MAX_RETRIES_OPENAI - 1: continue # Reintentar
                 else: break # Fallar tras último intento

            texto_respuesta_final = respuesta_inicial.choices[0].message.content.strip()
            finish_reason = respuesta_inicial.choices[0].finish_reason
            logger.info(f"Respuesta OpenAI OK (Longitud: {len(texto_respuesta_final)}, Razón Finalización: {finish_reason}).")
            if finish_reason == 'length':
                logger.warning(f"Respuesta OpenAI truncada por max_tokens ({respuesta_inicial.usage.completion_tokens if respuesta_inicial.usage else 'N/A'} tokens usados).")
                texto_respuesta_final += "\n<p><i>(Nota: La respuesta puede estar incompleta debido a límites de longitud.)</i></p>"


            # --- Búsqueda web si es necesario ---
            respuesta_lower = texto_respuesta_final.lower()
            # Revisar si la respuesta *explícitamente* dice no saber O si se forzó la búsqueda
            necesita_web = any(frase in respuesta_lower for frase in FRASES_BUSQUEDA) or forzar_busqueda_web

            if necesita_web:
                logger.info(f"Realizando búsqueda web (Forzado: {forzar_busqueda_web})...")
                web_resultados_html = buscar_google(mensaje_usuario)
                if web_resultados_html and not web_resultados_html.startswith("<p><i>["):
                     # Añadir resultados de forma elegante
                     texto_respuesta_final += "\n\n" + web_resultados_html
                     logger.info("Resultados de búsqueda web añadidos.")
                else:
                     logger.info("Búsqueda web no produjo resultados o encontró un error.")
            else:
                 logger.info("No se requiere búsqueda web para esta respuesta.")

            # Éxito, salir del bucle de reintentos
            break

         except APIError as e:
            logger.error(f"Error API OpenAI en /consulta (Intento {attempt + 1}): {e}", exc_info=True)
            texto_respuesta_final = f"<p><i>Error al contactar el servicio de IA: {e.message}. Inténtalo de nuevo más tarde.</i></p>"
            if attempt == MAX_RETRIES_OPENAI - 1:
                 # No lanzar HTTPException aquí para poder guardar historial con error
                 pass
            # time.sleep(1) # Espera opcional antes de reintentar
         except Exception as e:
            logger.error(f"Error inesperado en /consulta (Intento {attempt + 1}): {e}", exc_info=True)
            texto_respuesta_final = "<p><i>Ocurrió un error interno inesperado al procesar tu consulta.</i></p>"
            if attempt == MAX_RETRIES_OPENAI - 1:
                 pass
            # time.sleep(1)

    # --- Guardar en historial (incluso si hubo error OpenAI leve) ---
    if DB_CONFIGURED:
        conn_hist = get_db_connection()
        if conn_hist:
            try:
                with conn_hist.cursor() as cursor:
                    sql_hist = """
                        INSERT INTO historial (usuario_id, tenant_id, pregunta, respuesta, fecha_hora)
                        VALUES (%s, %s, %s, %s, NOW())
                    """
                    # Usar mensaje original y respuesta final (que puede contener mensaje de error)
                    cursor.execute(sql_hist, (current_user_id, current_tenant_id, mensaje_usuario, texto_respuesta_final))
                conn_hist.commit()
                logger.info(f"Consulta guardada en historial para user {current_user_id}/tenant {current_tenant_id}.")
            except (Exception, psycopg2.Error) as e_hist:
                logger.error(f"Error al guardar consulta en historial para user {current_user_id}: {e_hist}", exc_info=True)
                if conn_hist: conn_hist.rollback()
            finally:
                if conn_hist: conn_hist.close()
        else:
            logger.warning("No se pudo guardar consulta en historial (sin conexión BD).")


    # --- Limpieza final de la respuesta (opcional) ---
    # if BS4_AVAILABLE: ... (limpieza si es necesaria) ...

    return RespuestaConsulta(respuesta=texto_respuesta_final)


# --- Endpoint /analizar-documento ---
@app.post("/analizar-documento", response_model=RespuestaAnalisis)
async def analizar_documento(
    file: UploadFile = File(...),
    especializacion: str = Form("general"),
    user_id: int | None = Form(None),
    tenant_id: int | None = Form(None)
):
    """
    Analiza un documento subido (imagen o texto PDF/DOCX/TXT/CSV) usando OpenAI
    y devuelve un informe formateado en HTML.
    """
    if not client:
        logger.error("Llamada a /analizar-documento pero cliente OpenAI no está configurado.")
        raise HTTPException(status_code=503, detail="Servicio IA no disponible (Configuración).")

    # Validar IDs
    current_user_id = user_id
    current_tenant_id = tenant_id
    if not isinstance(current_user_id, int) or not isinstance(current_tenant_id, int):
        logger.error(f"IDs inválidos recibidos en /analizar-documento: User='{current_user_id}', Tenant='{current_tenant_id}'")
        raise HTTPException(status_code=400, detail="User ID y Tenant ID deben ser números enteros válidos.")

    # --- Mismo código que v2.4.0 para obtener prompt, procesar archivo, llamar a OpenAI ---
    filename = file.filename if file.filename else "archivo_subido"
    content_type = file.content_type or "application/octet-stream"
    # Obtener extensión de forma más segura
    base, dot, extension = filename.rpartition('.')
    extension = extension.lower() if dot else ''

    especializacion_lower = especializacion.lower() if especializacion else "general"
    logger.info(f"Solicitud Análisis Documento: User={current_user_id}, Tenant={current_tenant_id}, File='{filename}', Type='{content_type}', Espec='{especializacion_lower}'")

    # Obtener prompt personalizado (Memoria)
    custom_prompt_text = ""
    if DB_CONFIGURED: # Solo si BD está configurada
        conn_prompt = get_db_connection()
        if conn_prompt:
            try:
                with conn_prompt.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
                    sql_prompt = "SELECT custom_prompt FROM user_settings WHERE user_id = %s AND tenant_id = %s"
                    cursor.execute(sql_prompt, (current_user_id, current_tenant_id))
                    result = cursor.fetchone()
                    if result and result.get('custom_prompt') and result['custom_prompt'].strip():
                        custom_prompt_text = result['custom_prompt'].strip()
                        logger.info(f"Prompt personalizado (Memoria) encontrado para análisis (User {current_user_id}/Tenant {current_tenant_id}).")
            except (Exception, psycopg2.Error) as e_prompt:
                 logger.error(f"Error BD get prompt (Memoria) para análisis User {current_user_id}: {e_prompt}", exc_info=True)
            finally:
                conn_prompt.close()
        else:
            logger.warning(f"No conexión BD para prompt (Memoria) análisis.")


    # Construir System Prompt para Análisis
    prompt_especifico = PROMPT_ESPECIALIZACIONES.get(especializacion_lower, PROMPT_ESPECIALIZACIONES["general"])
    system_prompt_parts = [BASE_PROMPT_ANALISIS_DOC, prompt_especifico]
    if custom_prompt_text:
        system_prompt_parts.append(f"\n\n### Instrucciones Adicionales (Memoria) ###\n{custom_prompt_text}")
    system_prompt = "\n".join(filter(None, system_prompt_parts))
    # logger.debug(f"System Prompt Final para Análisis:\n------\n{system_prompt}\n------")

    # Preparar payload para OpenAI
    messages_payload = []
    IMAGE_MIMES = ["image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif"]
    TEXT_EXTENSIONS = ["pdf", "doc", "docx", "txt", "csv"]
    temp_filename_analisis = None # Para asegurar limpieza en finally

    try:
        if content_type in IMAGE_MIMES:
            logger.info(f"Procesando imagen '{filename}' ({content_type}) para análisis con GPT-4 Vision.")
            image_bytes = await file.read()
            # Validar tamaño máximo
            MAX_IMAGE_SIZE = 20 * 1024 * 1024 # OpenAI Vision limit (aprox)
            if len(image_bytes) > MAX_IMAGE_SIZE:
                 logger.error(f"Imagen '{filename}' excede el límite de {MAX_IMAGE_SIZE / (1024*1024):.1f} MB.")
                 raise HTTPException(status_code=413, detail=f"La imagen excede el tamaño máximo permitido ({MAX_IMAGE_SIZE / (1024*1024):.1f} MB).")

            base64_image = base64.b64encode(image_bytes).decode('utf-8')
            user_prompt = "Analiza detalladamente la siguiente imagen y genera un informe profesional y bien estructurado en formato HTML sobre su contenido, propósito y puntos clave."
            messages_payload = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": [
                    {"type": "text", "text": user_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{content_type};base64,{base64_image}"}}
                ]}
            ]
        elif extension in TEXT_EXTENSIONS:
            logger.info(f"Procesando archivo de texto '{filename}' ({extension.upper()}) para análisis.")
            texto_extraido = ""
            # Guardar temporalmente de forma segura usando NamedTemporaryFile
            with tempfile.NamedTemporaryFile(mode='wb', suffix=f'.{extension}', dir=TEMP_DIR, delete=False) as temp_file:
                 temp_filename_analisis = temp_file.name
                 try:
                     # Leer el archivo subido por chunks y escribir al temporal
                     while True:
                         chunk = await file.read(8192) # Leer en chunks
                         if not chunk:
                             break
                         temp_file.write(chunk)
                     logger.info(f"Archivo '{filename}' guardado temporalmente como '{temp_filename_analisis}'")
                 except Exception as copy_err:
                     logger.error(f"Error al copiar el archivo subido a temporal '{temp_filename_analisis}': {copy_err}", exc_info=True)
                     raise HTTPException(status_code=500, detail="Error al guardar el archivo temporalmente.")

            # Extraer texto del archivo temporal guardado (manejo de errores dentro de las funciones)
            try:
                 if extension in ['pdf', 'doc', 'docx']:
                    texto_extraido = extraer_texto_pdf_docx(temp_filename_analisis, extension)
                 else: # txt, csv
                    texto_extraido = extraer_texto_simple(temp_filename_analisis)
            finally:
                 # Borrar el archivo temporal después de la extracción SIEMPRE
                 if temp_filename_analisis and os.path.exists(temp_filename_analisis):
                     try:
                         os.remove(temp_filename_analisis)
                         logger.info(f"Archivo temporal '{temp_filename_analisis}' eliminado tras extracción.")
                     except OSError as e_remove:
                         logger.error(f"Error al eliminar archivo temporal '{temp_filename_analisis}': {e_remove}")
                     temp_filename_analisis = None # Marcar como borrado


            # Validar texto extraído antes de enviarlo a OpenAI
            if texto_extraido.startswith("[Error") or not texto_extraido.strip():
                error_msg = texto_extraido if texto_extraido.startswith("[Error") else "[Archivo vacío o sin texto extraíble]"
                logger.error(f"Error o texto vacío durante la extracción para '{filename}': {error_msg}")
                raise HTTPException(status_code=400, detail=f"Error al extraer texto del documento: {error_msg}")

            # Limitar longitud del texto enviado a OpenAI si es necesario (ej. ~100k tokens límite aprox)
            MAX_ANALYSIS_TOKENS = 100000
            estimated_tokens = len(texto_extraido.split()) * 1.3
            if estimated_tokens > MAX_ANALYSIS_TOKENS:
                 logger.warning(f"Texto extraído de '{filename}' excede el límite aprox. de {MAX_ANALYSIS_TOKENS} tokens. Se truncará.")
                 # Truncar por caracteres (menos preciso que por tokens)
                 max_chars = int(MAX_ANALYSIS_TOKENS / 1.3)
                 texto_extraido = texto_extraido[:max_chars] + "\n\n[... CONTENIDO TRUNCADO ...]"


            user_prompt = f"Redacta un informe HTML profesional y bien estructurado basado en el siguiente texto extraído del documento '{htmlspecialchars(filename)}':\n\n--- INICIO CONTENIDO DOCUMENTO ---\n{texto_extraido}\n--- FIN CONTENIDO DOCUMENTO ---"
            messages_payload = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
        else:
            # Tipo de archivo no soportado
            logger.error(f"Tipo de archivo no soportado para análisis: '{content_type or extension}' (Archivo: '{filename}')")
            raise HTTPException(status_code=415, detail=f"Tipo de archivo '{content_type or extension}' no soportado para análisis directo.")

        # Verificar que el payload se generó
        if not messages_payload:
             logger.critical("No se generó el payload para OpenAI en /analizar-documento. Error de lógica interna.")
             raise HTTPException(status_code=500, detail="Error interno: No se pudo preparar la solicitud para la IA.")

        # Llamada a OpenAI (con reintentos)
        informe_html = "<p><i>Error al generar el informe.</i></p>" # Default
        MAX_RETRIES_OPENAI = 2
        for attempt in range(MAX_RETRIES_OPENAI):
             try:
                 logger.info(f"Llamando a OpenAI para análisis de documento '{filename}' (Intento {attempt + 1}/{MAX_RETRIES_OPENAI})...")
                 respuesta_informe = client.chat.completions.create(
                     model="gpt-4-turbo", # Modelo adecuado para análisis de texto largo
                     messages=messages_payload,
                     temperature=0.4, # Ligeramente más determinista para informes
                     max_tokens=3000 # Límite generoso para el informe
                 )

                 if not respuesta_informe.choices or not respuesta_informe.choices[0].message or not respuesta_informe.choices[0].message.content:
                      logger.error(f"Respuesta de OpenAI inválida o vacía para análisis de '{filename}'.")
                      if attempt < MAX_RETRIES_OPENAI - 1: continue # Reintentar
                      else: raise ValueError("Respuesta inválida de OpenAI tras reintentos.") # Fallar


                 informe_html = respuesta_informe.choices[0].message.content.strip()
                 finish_reason = respuesta_informe.choices[0].finish_reason
                 logger.info(f"Informe generado por OpenAI para '{filename}' (Longitud: {len(informe_html)}, Razón Finalización: {finish_reason}).")
                 if finish_reason == 'length':
                    logger.warning(f"Informe OpenAI truncado por max_tokens para '{filename}'.")
                    informe_html += "\n<p><i>(Nota: El informe puede estar incompleto debido a límites de longitud.)</i></p>"

                 break # Salir si la llamada fue exitosa

             except APIError as e:
                 logger.error(f"Error API OpenAI en /analizar para '{filename}' (Intento {attempt + 1}): {e}", exc_info=True)
                 if attempt == MAX_RETRIES_OPENAI - 1:
                      raise HTTPException(status_code=503, detail=f"Error OpenAI al analizar: {e.message} (tras {MAX_RETRIES_OPENAI} intentos)")
                 # time.sleep(1)
             except Exception as e:
                 logger.error(f"Error inesperado llamando a OpenAI en /analizar para '{filename}' (Intento {attempt + 1}): {e}", exc_info=True)
                 if attempt == MAX_RETRIES_OPENAI - 1:
                      raise HTTPException(status_code=500, detail=f"Error interno durante el análisis ({type(e).__name__})")
                 # time.sleep(1)


        # Limpieza básica del HTML generado (opcional pero recomendado)
        if BS4_AVAILABLE:
            try:
                # Quitar doctype, html, head si OpenAI los añade por error
                if "<!DOCTYPE html>" in informe_html or "<html" in informe_html:
                    soup = BeautifulSoup(informe_html, 'html.parser')
                    if soup.body:
                         informe_html = soup.body.decode_contents()
                         logger.info("Detectado HTML completo en informe, se ha extraído solo el contenido del body.")
                    elif soup.html:
                         informe_html = soup.html.decode_contents()
                         logger.info("Detectado tag <html> en informe, se ha extraído su contenido.")

                # Quitar ```html ... ``` si los añade
                informe_html = re.sub(r'^```[a-zA-Z]*\s*', '', informe_html, flags=re.IGNORECASE)
                informe_html = re.sub(r'\s*```$', '', informe_html)
                informe_html = informe_html.strip()

            except Exception as e_bs4:
                logger.error(f"Error al procesar/limpiar HTML del informe con BeautifulSoup: {e_bs4}")
        # Asegurar que la respuesta sea al menos un párrafo si no es HTML válido
        if not re.search(r'<[a-z][\s\S]*>', informe_html, re.IGNORECASE):
            logger.warning("La respuesta de análisis de OpenAI no parece ser HTML válido, envolviendo en <p>.")
            informe_html = f"<p>{htmlspecialchars(informe_html)}</p>" # Escapar por si acaso

        return RespuestaAnalisis(informe=informe_html)

    # Captura general de errores del endpoint /analizar-documento
    except HTTPException as e:
        # Re-lanzar excepciones HTTP ya generadas (ej. 400, 413, 415)
        raise e
    except Exception as e:
        logger.error(f"Error general inesperado en /analizar-documento para archivo '{filename}': {e}", exc_info=True)
        # Borrar archivo temporal si aún existe y ocurrió un error antes de borrarlo normalmente
        if temp_filename_analisis and os.path.exists(temp_filename_analisis):
            try:
                os.remove(temp_filename_analisis)
                logger.info(f"Archivo temporal '{temp_filename_analisis}' eliminado en bloque catch general.")
            except OSError as e_final_remove:
                logger.error(f"Error al eliminar archivo temporal '{temp_filename_analisis}' en catch general: {e_final_remove}")
        raise HTTPException(status_code=500, detail=f"Error interno del servidor al procesar el archivo ({type(e).__name__}).")
    finally:
         # Asegurarse de cerrar el archivo subido en caso de error también
         if file and not file.file.closed:
             await file.close()


# --- Endpoint /direccion/detalles/{place_id} ---
@app.get("/direccion/detalles/{place_id}", response_model=PlaceDetailsResponse)
async def obtener_detalles_direccion(
    place_id: str = Path(..., description="ID del lugar obtenido de Google Places Autocomplete"),
    user_id: int | None = Query(None, description="ID del usuario que realiza la solicitud (opcional para logging)"),
    tenant_id: int | None = Query(None, description="ID del tenant del usuario (opcional para logging)")
):
    """
    Obtiene detalles de una dirección (código postal, localidad, provincia, país)
    a partir de un Place ID de Google Places.
    Utiliza la API Key segura (`MAPS_API_ALL`) configurada en el backend.
    """
    logger.info(f"Solicitud detalles dirección para Place ID: {place_id} (User: {user_id}, Tenant: {tenant_id})")

    if not MAPS_CONFIGURED:
        logger.error("Endpoint /direccion/detalles llamado pero MAPS_API_ALL no está configurada.")
        # Devolver error 503 Service Unavailable si la API no está lista
        raise HTTPException(status_code=503, detail="Servicio de direcciones no disponible (configuración backend).")

    GOOGLE_PLACES_DETAILS_URL = "https://maps.googleapis.com/maps/api/place/details/json"
    # Campos necesarios para extraer CP, Localidad, Provincia, País
    fields_needed = "address_component,formatted_address" # Pedir componentes y dirección formateada

    params = {
        "place_id": place_id,
        "key": MAPS_API_ALL,
        "fields": fields_needed,
        "language": "es" # Preferir resultados en español
    }

    try:
        # Usar httpx para llamada asíncrona no bloqueante
        async with httpx.AsyncClient(timeout=10.0) as client_http: # Timeout de 10 segundos
            logger.debug(f"Llamando a Google Places Details API. URL: {GOOGLE_PLACES_DETAILS_URL}, Params: {params}")
            response = await client_http.get(GOOGLE_PLACES_DETAILS_URL, params=params)
            response.raise_for_status() # Lanza excepción para errores HTTP 4xx/5xx

            data = response.json()
            logger.debug(f"Respuesta JSON de Google Places API: {data}")

            # Verificar el estado devuelto por la API de Google
            api_status = data.get("status")
            if api_status != "OK":
                error_message = data.get("error_message", f"Estado API Google Places no OK: {api_status}")
                logger.error(f"Error API Google Places para place_id {place_id}: {error_message} (Status: {api_status})")
                # Mapear status a errores HTTP o devolver respuesta controlada
                if api_status == "REQUEST_DENIED":
                     # Podría ser problema de API Key inválida, no habilitada, o restricciones
                     raise HTTPException(status_code=403, detail="Acceso denegado por Google Places API. Verifica la clave API del backend y sus permisos.")
                elif api_status == "INVALID_REQUEST":
                      raise HTTPException(status_code=400, detail="Solicitud inválida a Google Places API (place_id incorrecto?).")
                elif api_status == "ZERO_RESULTS":
                      logger.warning(f"Google Places API no encontró detalles para place_id {place_id}.")
                      return PlaceDetailsResponse(success=False, error="No se encontraron detalles para la dirección seleccionada.")
                else: # OVER_QUERY_LIMIT, NOT_FOUND, UNKNOWN_ERROR
                     raise HTTPException(status_code=503, detail=f"Google Places API devolvió: {api_status}")

            # Procesar resultado si status es OK
            result = data.get("result", {})
            address_components = result.get("address_components", [])
            formatted_address = result.get("formatted_address", "N/A")
            logger.info(f"Dirección formateada recibida para {place_id}: {formatted_address}")

            if not address_components:
                 logger.warning(f"No se encontraron 'address_components' en la respuesta OK de Google para {place_id}")
                 # Devolver éxito pero sin datos, o fallo? Fallo parece más claro.
                 return PlaceDetailsResponse(success=False, error="No se recibieron componentes de dirección detallados.")

            # Extraer los datos necesarios
            postal_code = None
            locality = None
            province = None # Variable para provincia
            country = None

            for component in address_components:
                types = component.get("types", [])
                long_name = component.get("long_name")
                if not types or not long_name: continue # Saltar componente inválido

                # Extraer componentes basado en 'types'
                if "postal_code" in types: postal_code = long_name
                if "locality" in types: locality = long_name
                # Provincia/Estado (puede variar según país, priorizamos level_2)
                if "administrative_area_level_2" in types: province = long_name
                elif "administrative_area_level_1" in types and not province: province = long_name # Fallback a nivel 1
                if "country" in types: country = long_name

                # Lógica adicional si 'locality' falta (menos común con types=address)
                # if not locality and "administrative_area_level_3" in types: locality = long_name
                # elif not locality and province and "administrative_area_level_2" in types: locality = province # Usar provincia como localidad?

            logger.info(f"Componentes extraídos para {place_id}: CP='{postal_code}', Loc='{locality}', Prov='{province}', Pais='{country}'")

            # Devolver respuesta exitosa con los datos encontrados (pueden ser None)
            return PlaceDetailsResponse(
                success=True,
                postal_code=postal_code,
                locality=locality,
                province=province,
                country=country
            )

    # Manejo de errores de la llamada HTTP o procesamiento
    except httpx.TimeoutException:
        logger.error(f"Timeout al contactar Google Places API para place_id {place_id}")
        raise HTTPException(status_code=504, detail="Timeout al obtener detalles de la dirección.")
    except httpx.RequestError as e:
        logger.error(f"Error de conexión al contactar Google Places API para place_id {place_id}: {e}", exc_info=True)
        raise HTTPException(status_code=503, detail="Error de conexión al obtener detalles de la dirección.")
    except HTTPException as e: # Re-lanzar excepciones HTTP ya generadas (ej. 400, 403, 503 de Google)
        raise e
    except Exception as e: # Otros errores inesperados (ej. JSON parsing, etc.)
        logger.error(f"Error inesperado procesando detalles de dirección para place_id {place_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Error interno del servidor al procesar la dirección.")

# --- FIN NUEVO ENDPOINT ---


# --- Punto de Entrada para ejecución local (Uvicorn) ---
# Descomentar la línea si se quiere ejecutar localmente con `python main.py`
# Asegúrate de tener un archivo .env con las variables de entorno si ejecutas localmente.
# if __name__ == "__main__":
#     import uvicorn
#     logger.info("Iniciando servidor Uvicorn para desarrollo local...")
#     # Cargar variables .env si se usa python-dotenv (necesitarías añadirlo a requirements)
#     # from dotenv import load_dotenv
#     # load_dotenv()
#     uvicorn.run("main:app", host="0.0.0.0", port=int(os.getenv("PORT", 8000)), reload=True)

# --- FIN main.py v2.4.1-mt ---